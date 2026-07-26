"""
livrables.py — Rapport station « prêt à copier » (package thermie_debits).

Produit deux livrables harmonisés pour une station donnée :

  • un rapport Word (.docx) organisé en chapitres séparés par des sauts de
    page, axé résultats — les rappels méthodologiques y sont réduits à de
    courts et centrés sur les résultats ;
  • un classeur Excel (.xlsx) organisé en onglets miroirs du rapport, dont
    le dernier porte une LIGNE DE SYNTHÈSE À SCHÉMA FIXE permettant
    d'empiler les stations les unes sous les autres sans retraitement.

Les textes d'interprétation proviennent tous de `redaction.py` (source
unique partagée avec l'application) : le rapport dit nécessairement la même
chose que ce que l'utilisateur a lu à l'écran.
"""
from __future__ import annotations
import io
from datetime import date

import numpy as np
import pandas as pd
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT

from . import redaction as red
from .print_style import make_print_ready
from .config import __version__, VERSION_DATE

# Palette partagée docx / xlsx
COUL = {"ok": "1E8449", "attention": "B9770D", "alerte": "C0392B",
        "neutre": "5D6D7E"}
BLEU = "1A5276"
GRIS = "5D6D7E"

LARGEUR_PORTRAIT_CM = 16.0


# ============================================================
# OUTILS DOCX
# ============================================================
def _titre_chapitre(doc, num, texte):
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run(f"{num}. {texte}")
    r.bold = True; r.font.size = Pt(17); r.font.color.rgb = RGBColor.from_string(BLEU)
    p.space_after = Pt(10)
    return p


def _sous_titre(doc, texte):
    p = doc.add_paragraph()
    r = p.add_run(texte)
    r.bold = True; r.font.size = Pt(12.5)
    r.font.color.rgb = RGBColor.from_string("1A1A2E")
    p.space_before = Pt(12); p.space_after = Pt(5)
    return p


def _para(doc, texte, italique=False, taille=10.5, couleur=None):
    p = doc.add_paragraph()
    r = p.add_run(texte)
    r.italic = italique; r.font.size = Pt(taille)
    if couleur:
        r.font.color.rgb = RGBColor.from_string(couleur)
    p.space_after = Pt(5)
    return p


def _encadre(doc, bloc):
    """Encadré d'interprétation (titre + verdict coloré + lignes)."""
    coul = COUL.get(bloc.get("niveau", "neutre"), GRIS)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(LARGEUR_PORTRAIT_CM)

    p0 = cell.paragraphs[0]
    r0 = p0.add_run(bloc["titre"])
    r0.bold = True; r0.font.size = Pt(11); r0.font.color.rgb = RGBColor.from_string(coul)

    p1 = cell.add_paragraph()
    r1 = p1.add_run(f"→ {bloc['verdict']}")
    r1.bold = True; r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor.from_string(coul)

    for l in bloc.get("lignes", []):
        pl = cell.add_paragraph()
        rl = pl.add_run("• " + str(l))
        rl.font.size = Pt(9.5)
        pl.space_after = Pt(2)
    doc.add_paragraph()
    return tbl


def _figure(doc, fig, legende=None, largeur_cm=LARGEUR_PORTRAIT_CM):
    """Insère une figure matplotlib à la largeur d'impression cible.
    La copie « prête à imprimer » garantit une police ≥ 8 pt une fois
    l'image réduite à cette largeur (voir print_style)."""
    if fig is None:
        return
    buf = io.BytesIO()
    fig_impr = make_print_ready(fig)
    fig_impr.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    doc.add_picture(buf, width=Cm(largeur_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if legende:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(legende)
        r.italic = True; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(GRIS)


def _tableau(doc, entetes, lignes, largeurs=None):
    if not lignes:
        return
    tbl = doc.add_table(rows=1, cols=len(entetes))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, e in enumerate(entetes):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(str(e))
        r.bold = True; r.font.size = Pt(9)
    for lg in lignes:
        cells = tbl.add_row().cells
        for i, v in enumerate(lg):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run("" if v is None else str(v))
            r.font.size = Pt(8.5)
    if largeurs:
        for row in tbl.rows:
            for i, w in enumerate(largeurs):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return tbl


# ============================================================
# RAPPORT DOCX
# ============================================================
def construire_docx_bytes(res):
    doc = Document()
    for s in doc.sections:
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin = s.right_margin = Cm(2.5)
        s.top_margin = s.bottom_margin = Cm(2.0)

    cfg = res.config.sources
    nom = cfg.nom_cours_eau or "Station"
    F = res.figures or {}

    # ---------- Page de garde ----------
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("APPROCHE THERMIQUE")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor.from_string(BLEU)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Rapport de station")
    r.font.size = Pt(16); r.font.color.rgb = RGBColor.from_string(GRIS)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(nom); r.bold = True; r.font.size = Pt(22)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(res.contexte["label"]); r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(GRIS)
    for _ in range(6):
        doc.add_paragraph()

    infos = [("Localisation de la sonde", getattr(cfg, "localisation_sonde", "") or "—"),
             ("Station hydrométrique", getattr(cfg, "nom_station_debit", "") or "—"),
             ("Mode d'analyse", "thermie + débits" if res.config.avec_debits
              else "thermie seule"),
             ("Date de production", f"{date.today():%d/%m/%Y}"),
             ("Version de l'outil", f"{__version__} ({VERSION_DATE})")]
    _tableau(doc, ["Paramètre", "Valeur"], infos, largeurs=[7.0, 9.0])

    # ---------- Chapitre 1 — Données et qualité ----------
    _titre_chapitre(doc, 1, "Données mobilisées et qualité")
    _para(doc, "Ce chapitre établit ce sur quoi repose l'analyse : quelles "
               "données ont été chargées, ce que le contrôle qualité a écarté, "
               "et — point déterminant — si le modèle de transfert air–eau est "
               "suffisamment robuste pour que la normalisation climatique, dont "
               "dépend l'ensemble des indicateurs, soit fiable.")

    _sous_titre(doc, "1.1 Sources chargées")
    _tableau(doc, ["Source", "Contenu"], red.synthese_sources(res),
             largeurs=[5.5, 10.5])

    _sous_titre(doc, "1.2 Contrôle qualité")
    _encadre(doc, red.synthese_qc(res))
    _figure(doc, F.get("qc"), "Chronique brute et enregistrements écartés par le "
                              "contrôle qualité.")

    _sous_titre(doc, "1.3 Couverture calendaire")
    _encadre(doc, red.couverture_calendaire(res))

    _sous_titre(doc, "1.4 Chronique thermique")
    _figure(doc, F.get("chronique"), "Chronique thermique : température de l'air, "
                                     "normales 1991–2020, et température de l'eau "
                                     "brute puis compensée.")

    _sous_titre(doc, "1.5 Validité du modèle air–eau")
    _para(doc, "L'intérêt de cette normalisation est de rendre les stations et "
               "les années comparables entre elles, malgré des contextes "
               "physiques et des périodes d'enregistrement hétérogènes : sans "
               "elle, une année plus chaude ou une chronique plus courte "
               "biaiserait directement la lecture de la vulnérabilité.")
    _para(doc, "La normalisation climatique retranche à la température observée "
               "l'anomalie de l'air pondérée par le coefficient de sensibilité. "
               "Presque tous les indicateurs de la méthode en découlent : la "
               "validité de cette régression conditionne donc la portée de "
               "l'ensemble des résultats qui suivent.")
    _encadre(doc, red.validite_modele_air_eau(res))
    _figure(doc, F.get("sensibilite"), "Régression température de l'air / "
                                       "température de l'eau (juin–septembre) et "
                                       "indicateurs de validité.")

    if res.figures_climatiques:
        _sous_titre(doc, "1.6 Contexte météorologique et climatique")
        _para(doc, "Volet descriptif sur données brutes, distinct des volets "
                   "analytiques (qui appliquent le contrôle qualité). Il situe la "
                   "période de mesure dans son contexte pluriannuel.")
        for f in res.figures_climatiques:
            _figure(doc, f, f.axes[0].get_title().split("\n")[0])

    # ---------- Chapitre 2 — Thermie estivale ----------
    _titre_chapitre(doc, 2, "Vulnérabilité thermique estivale")
    _para(doc, "Caractérisation de la contrainte thermique subie en période "
               "estivale (juin–septembre), indépendamment de toute donnée de "
               "débit : stress métabolique cumulé d'une part, risque de "
               "mortalité à court terme d'autre part.")

    _sous_titre(doc, "2.1 Vulnérabilité chronique et aiguë")
    _encadre(doc, red.interpretation_vulnerabilite(res))
    _figure(doc, F.get("vulnerabilite"), "Températures normalisées et seuils de "
                                         "stress et de létalité, une couleur par "
                                         "année.")

    _sous_titre(doc, "2.2 Indicateurs thermiques descriptifs")
    _figure(doc, F.get("indicateurs_resume"), "Températures extrêmes et amplitude "
                                              "nycthémérale par mois.")
    if res.indicateurs is not None:
        t = res.indicateurs.get("table_mensuelle")
        if t is not None and len(t):
            cols = [c for c in t.columns][:6]
            lignes = [[("" if pd.isna(v) else v) for v in row[:6]]
                      for row in t[cols].values.tolist()]
            _tableau(doc, cols, lignes)
        # Corrélations n'impliquant pas le débit — toujours calculables,
        # y compris sur une station sans hydrométrie. Les deux corrélations
        # avec débit sont présentées au chapitre 4 (relation débit–température).
        from . import figures as figmod
        cor = res.indicateurs.get("correlations") or {}
        fig_cor_sans_q = figmod.fig_correlations_indicateurs(
            cor, res.config.sources.nom_cours_eau, None,
            cles={"ampl_vs_teau", "ecart_vs_teau"},
            suffixe_titre="Indicateurs vs température de l'eau",
            filename="Fig_Correlations_sans_Debit.png")
        _figure(doc, fig_cor_sans_q, "Amplitude nycthémérale et écart eau–air en "
                                     "fonction de la température de l'eau.")

    # ---------- Chapitre 3 — Fraie-croissance ----------
    fr = res.fraie
    if fr is not None:
        _titre_chapitre(doc, 3, "Vulnérabilité de la reproduction")
        _para(doc, "Besoins thermiques propres à la reproduction, évalués sur "
                   "trois phases successives (pré-frai, ponte, incubation) aux "
                   "tolérances distinctes. Ce volet couvre une période de "
                   "l'année différente de l'étiage estival.")
        _encadre(doc, red.interpretation_fraie(res))
        _figure(doc, F.get("fraie"), "Températures normalisées confrontées aux "
                                     "fenêtres thermiques de chaque phase, une "
                                     "couleur par campagne annuelle.")
        lignes = []
        for s in fr.get("sous_indicateurs", []):
            for ph in s.get("phases", []):
                if not ph.get("n"):
                    continue
                lignes.append([s["espece"], ph["nom"],
                               f"{ph['opt'][0]:.0f}–{ph['opt'][1]:.0f}",
                               f"{ph['elargie'][0]:.0f}–{ph['elargie'][1]:.0f}",
                               ph["n"], f"{ph['pct_optimum']:.0f}",
                               f"{ph['pct_elargie']:.0f}", f"{ph['pct_letal']:.0f}"])
        _tableau(doc, ["Espèce", "Phase", "Optimum °C", "Élargie °C", "n j",
                       "% opt.", "% élar.", "% létal"], lignes,
                 largeurs=[2.4, 3.4, 2.0, 2.0, 1.3, 1.6, 1.6, 1.6])

    # ---------- Chapitre 4 — Relation Q–T° et débits ----------
    num_ch4 = 4 if fr is not None else 3
    if res.config.avec_debits:
        _titre_chapitre(doc, num_ch4, "Relation débit–température et débits de "
                                      "référence")
        _para(doc, "Toute la démarche des débits de référence thermique repose "
                   "sur un postulat : le débit module la température de l'eau. "
                   "Ce chapitre le teste explicitement, en tire un diagnostic sur "
                   "la validité de la démarche pour cette station, puis présente "
                   "les débits qui en découlent.")

        _sous_titre(doc, f"{num_ch4}.1 Le débit module-t-il la température ?")
        _encadre(doc, red.interpretation_relation(res))
        _figure(doc, F.get("relation_debit_temp"), "Relation observée et effet "
                                                   "propre du débit, à température "
                                                   "d'air égale.")

        _sous_titre(doc, f"{num_ch4}.2 Diagnostic — validité de la démarche")
        _encadre(doc, red.interpretation_matrice(res))
        _figure(doc, F.get("matrice"), "Croisement du constat de vulnérabilité et "
                                       "du test de relation débit–température.")

        _sous_titre(doc, f"{num_ch4}.3 Indicateurs dépendant du débit")
        from . import figures as figmod
        cor = (res.indicateurs or {}).get("correlations") or {}
        fig_cor_avec_q = figmod.fig_correlations_indicateurs(
            cor, res.config.sources.nom_cours_eau, None,
            cles={"ampl_vs_debit", "ecart_vs_debit"},
            suffixe_titre="Indicateurs vs débit (échelle logarithmique)",
            filename="Fig_Correlations_avec_Debit.png")
        _figure(doc, fig_cor_avec_q, "Amplitude nycthémérale et écart eau–air en "
                                     "fonction du débit.")

        _sous_titre(doc, f"{num_ch4}.4 Débits de référence thermique")
        _encadre(doc, red.interpretation_debits(res))
        for k, leg in [("debits_vuln", "Vulnérabilité biologique en fonction du "
                                       "débit — base de Q_thermie_bio."),
                       ("debits_inflexion", "Rupture de régime thermique — base de "
                                            "Q_thermie_fonc (appoint)."),
                       ("debits_classes", "Positionnement des débits de référence "
                                          "sur les courbes de débits classés.")]:
            _figure(doc, F.get(k), leg)

    # ---------- Chapitre final — Synthèse ----------
    num_fin = num_ch4 + (1 if res.config.avec_debits else 0)
    _titre_chapitre(doc, num_fin, "Synthèse et conclusions")

    if res.config.avec_debits:
        _sous_titre(doc, "Débit de référence retenu")
        _para(doc, "Point de sortie majeur de l'étude : le débit ci-dessous est "
                   "celui à retenir pour l'arbitrage, mis en avant avant le "
                   "détail des composantes qui suivent.")
        _encadre(doc, red.debit_retenu_prominent(res))

    _sous_titre(doc, "Score global de vulnérabilité thermique")
    _para(doc, "Le SGVT agrège les composantes établies dans les chapitres "
               "précédents. Il est présenté ici, une fois toutes ses composantes "
               "connues, et constitue une information d'appoint contextualisant "
               "les débits de référence.")
    _encadre(doc, red.interpretation_sgvt(res))
    _figure(doc, F.get("synthese_jauge"), "Positionnement du SGVT sur la grille "
                                          "de risque.", largeur_cm=11.0)
    _figure(doc, F.get("synthese_tableau"), "Détail des composantes du score.")

    _sous_titre(doc, "Conclusion")
    _para(doc, red.conclusion_station(res))

    _sous_titre(doc, "Réserves et limites propres à cette station")
    for r_ in red.reserves_station(res):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(r_); run.font.size = Pt(10)

    _sous_titre(doc, "Fiche de synthèse (schéma commun à toutes les stations)")
    _para(doc, "Ces valeurs constituent la ligne de synthèse de la station, "
               "reprise à l'identique dans l'onglet « SYNTHÈSE » du classeur "
               "Excel associé, afin de permettre la comparaison inter-sites.",
          italique=True, taille=9, couleur=GRIS)
    ls = red.ligne_synthese(res)
    _tableau(doc, ["Indicateur", "Valeur"],
             [[k, v] for k, v in ls.items()], largeurs=[7.5, 8.5])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# CLASSEUR XLSX
# ============================================================
_THIN = Side(style="thin", color="BFC9CA")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)


def _xl_titre(ws, ligne, texte, couleur=BLEU, taille=13):
    c = ws.cell(row=ligne, column=1, value=texte)
    c.font = Font(bold=True, size=taille, color=couleur)
    return ligne + 2


def _xl_bloc(ws, ligne, bloc):
    """Encadré d'interprétation dans une feuille Excel."""
    coul = COUL.get(bloc.get("niveau", "neutre"), GRIS)
    c = ws.cell(row=ligne, column=1, value=bloc["titre"])
    c.font = Font(bold=True, size=11, color=coul); ligne += 1
    c = ws.cell(row=ligne, column=1, value=f"→ {bloc['verdict']}")
    c.font = Font(bold=True, size=10.5, color=coul); ligne += 1
    for l in bloc.get("lignes", []):
        c = ws.cell(row=ligne, column=1, value="• " + str(l))
        c.font = Font(size=9.5); c.alignment = Alignment(wrap_text=True, vertical="top")
        ws.row_dimensions[ligne].height = max(15, 13 * (1 + len(str(l)) // 110))
        ligne += 1
    return ligne + 1


def _xl_tableau(ws, ligne, entetes, lignes, largeurs=None):
    fill = PatternFill("solid", fgColor=BLEU)
    for i, e in enumerate(entetes, start=1):
        c = ws.cell(row=ligne, column=i, value=str(e))
        c.font = Font(bold=True, color="FFFFFF", size=10)
        c.fill = fill; c.border = _BORDER
        c.alignment = Alignment(wrap_text=True, vertical="center")
        if largeurs and i <= len(largeurs):
            ws.column_dimensions[get_column_letter(i)].width = largeurs[i - 1]
    ligne += 1
    for lg in lignes:
        for i, v in enumerate(lg, start=1):
            c = ws.cell(row=ligne, column=i,
                        value=("" if v is None or (isinstance(v, float) and np.isnan(v))
                               else v))
            c.font = Font(size=9.5); c.border = _BORDER
            c.alignment = Alignment(wrap_text=True, vertical="top")
        ligne += 1
    return ligne + 1


def construire_xlsx_bytes(res):
    wb = openpyxl.Workbook()
    cfg = res.config.sources

    # ---- 00 Fiche station ----
    ws = wb.active; ws.title = "00_Fiche station"
    ws.column_dimensions["A"].width = 42
    ws.column_dimensions["B"].width = 60
    l = _xl_titre(ws, 1, f"{cfg.nom_cours_eau or 'Station'} — Approche thermique")
    l = _xl_tableau(ws, l, ["Paramètre", "Valeur"], [
        ["Cours d'eau", cfg.nom_cours_eau or "—"],
        ["Localisation sonde", getattr(cfg, "localisation_sonde", "") or "—"],
        ["Station hydrométrique", getattr(cfg, "nom_station_debit", "") or "—"],
        ["Contexte piscicole", res.contexte["label"]],
        ["Mode d'analyse", "thermie + débits" if res.config.avec_debits
         else "thermie seule"],
        ["Date de production", f"{date.today():%d/%m/%Y}"],
        ["Version de l'outil", f"{__version__} ({VERSION_DATE})"],
    ], largeurs=[42, 60])
    if res.config.avec_debits:
        l = _xl_bloc(ws, l, red.debit_retenu_prominent(res))
    l = _xl_bloc(ws, l, red.interpretation_matrice(res))
    ws.cell(row=l, column=1, value="Conclusion").font = Font(bold=True, size=11)
    c = ws.cell(row=l + 1, column=1, value=red.conclusion_station(res))
    c.alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[l + 1].height = 70

    # ---- 01 Données & QC ----
    ws = wb.create_sheet("01_Donnees_QC")
    ws.column_dimensions["A"].width = 44; ws.column_dimensions["B"].width = 70
    l = _xl_titre(ws, 1, "Données mobilisées et qualité")
    l = _xl_tableau(ws, l, ["Source", "Contenu"], red.synthese_sources(res),
                    largeurs=[44, 70])
    l = _xl_bloc(ws, l, red.synthese_qc(res))
    l = _xl_bloc(ws, l, red.couverture_calendaire(res))
    l = _xl_bloc(ws, l, red.validite_modele_air_eau(res))
    if res.rapport_qc is not None and len(res.rapport_qc):
        l = _xl_titre(ws, l, "Détail des enregistrements écartés", GRIS, 11)
        rap = res.rapport_qc.head(300)
        l = _xl_tableau(ws, l, list(rap.columns),
                        rap.values.tolist(), largeurs=[18] * len(rap.columns))

    # ---- 02 Thermie & SGVT ----
    ws = wb.create_sheet("02_Thermie_SGVT")
    ws.column_dimensions["A"].width = 44; ws.column_dimensions["B"].width = 40
    l = _xl_titre(ws, 1, "Vulnérabilité thermique estivale et SGVT")
    l = _xl_bloc(ws, l, red.interpretation_vulnerabilite(res))
    s, v, sg = res.sensibilite, res.vulnerabilite, res.sgvt
    if s and v and sg:
        l = _xl_tableau(ws, l, ["Indicateur", "Valeur", "Catégorie"], [
            ["Pente m (sensibilité)", round(s["m"], 4), s["sens_cat"]],
            ["R² air–eau", round(s["r2"], 4), s["r2_cat"]],
            ["Indice robustesse |ρ−r|", round(s["robustesse"], 4), s["rob_cat"]],
            [f"Stress chronique (Tmh > {v['seuil_chr']}°C)",
             round(v["pct_chr"], 2), v["cat_chr"]],
            [f"Létalité aiguë (Tmax > {v['seuil_aigu']}°C)",
             int(v["n_aigu"]), v["cat_aigu"]],
            ["SGVT /10", round(sg["sgvt"], 2), sg["interp"]],
            ["Composantes SGVT", sg.get("composantes", "—"), ""],
        ], largeurs=[44, 20, 40])
    l = _xl_bloc(ws, l, red.interpretation_sgvt(res))

    # ---- 03 Fraie ----
    if res.fraie is not None:
        ws = wb.create_sheet("03_Fraie")
        ws.column_dimensions["A"].width = 22; ws.column_dimensions["B"].width = 30
        l = _xl_titre(ws, 1, "Vulnérabilité de la reproduction")
        l = _xl_bloc(ws, l, red.interpretation_fraie(res))
        lignes = []
        for si in res.fraie.get("sous_indicateurs", []):
            for ph in si.get("phases", []):
                if not ph.get("n"):
                    lignes.append([si["espece"], ph["nom"], "aucune donnée",
                                   "", 0, "", "", "", ""])
                    continue
                lignes.append([
                    si["espece"], ph["nom"],
                    f"{ph['opt'][0]:.0f}–{ph['opt'][1]:.0f} °C",
                    f"{ph['elargie'][0]:.0f}–{ph['elargie'][1]:.0f} °C",
                    ph["n"], round(ph["pct_optimum"], 1),
                    round(ph["pct_elargie"], 1), round(ph["pct_letal"], 1),
                    "oui" if ph.get("critique") else "non"])
        l = _xl_tableau(ws, l, ["Espèce", "Phase", "Optimum", "Élargie", "n jours",
                                "% optimum", "% élargie", "% létal", "Critique"],
                        lignes, largeurs=[22, 30, 16, 16, 10, 11, 11, 10, 10])

    # ---- 04 Débits & relation ----
    if res.config.avec_debits:
        ws = wb.create_sheet("04_Debits_Relation")
        ws.column_dimensions["A"].width = 44; ws.column_dimensions["B"].width = 40
        l = _xl_titre(ws, 1, "Relation débit–température et débits de référence")
        l = _xl_bloc(ws, l, red.interpretation_relation(res))
        rel = res.relation_debit_temp
        if rel and rel.get("disponible"):
            l = _xl_tableau(ws, l, ["Gamme de débit", "n", "r brute", "R² brute",
                                    "r partielle", "R² partielle", "Concluante"],
                            [[g["gamme"], g["n"], round(g["r_brute"], 4),
                              round(g["r2_brute"], 4), round(g["r_partielle"], 4),
                              round(g["r2_partielle"], 4),
                              "oui" if g["concluante"] else "non"]
                             for g in rel["lignes"]],
                            largeurs=[26, 8, 11, 11, 12, 12, 11])
        l = _xl_bloc(ws, l, red.interpretation_matrice(res))
        l = _xl_bloc(ws, l, red.interpretation_debits(res))
        ds = res.debits_sorties or {}
        bio, fonc = ds.get("q_thermie_bio") or {}, ds.get("q_thermie_fonc") or {}
        l = _xl_tableau(ws, l, ["Débit de référence", "Valeur (m³/s)",
                                "PNDA désinfluencé (%)", "PNDA influencé (%)"], [
            ["Q_thermie_bio (résultat principal)",
             bio.get("valeur"), bio.get("pnda_desinf"), bio.get("pnda_inf")],
            ["Q_thermie_fonc (appoint)",
             fonc.get("valeur"), fonc.get("pnda_desinf"), fonc.get("pnda_inf")],
        ], largeurs=[40, 16, 22, 20])

    # ---- 05 Indicateurs ----
    if res.indicateurs is not None:
        ws = wb.create_sheet("05_Indicateurs")
        l = _xl_titre(ws, 1, "Indicateurs thermiques mensuels et annuels")
        t = res.indicateurs.get("table_mensuelle")
        if t is not None and len(t):
            l = _xl_tableau(ws, l, list(t.columns), t.values.tolist(),
                            largeurs=[16] * len(t.columns))
        cor = res.indicateurs.get("correlations") or {}
        lignes = [[k, c.get("xlabel", ""), c.get("ylabel", ""), c.get("n", 0),
                   round(c["r2"], 4) if c.get("r2") == c.get("r2") else "",
                   round(c["pente"], 4) if c.get("pente") == c.get("pente") else "",
                   "log(Q)" if c.get("log_x") else "linéaire"]
                  for k, c in cor.items()]
        if lignes:
            l = _xl_tableau(ws, l, ["Corrélation", "Abscisse", "Ordonnée", "n",
                                    "R²", "Pente", "Régression sur"],
                            lignes, largeurs=[20, 24, 28, 8, 10, 10, 14])

    # ---- 06 SYNTHÈSE (schéma fixe) ----
    ws = wb.create_sheet("06_SYNTHESE")
    l = _xl_titre(ws, 1, "Fiche de synthèse — schéma commun à toutes les stations")
    c = ws.cell(row=2, column=1,
                value="Colonnes invariantes : une station sans débit produit les "
                      "mêmes colonnes, renseignées « non applicable ». Les lignes "
                      "de plusieurs stations peuvent donc être empilées telles "
                      "quelles pour la comparaison inter-sites.")
    c.font = Font(italic=True, size=9, color=GRIS)
    c.alignment = Alignment(wrap_text=True)
    ws.row_dimensions[2].height = 30

    ligne = red.ligne_synthese(res)
    fill = PatternFill("solid", fgColor=BLEU)
    for i, col in enumerate(red.COLONNES_SYNTHESE, start=1):
        c = ws.cell(row=4, column=i, value=col)
        c.font = Font(bold=True, color="FFFFFF", size=9.5)
        c.fill = fill; c.border = _BORDER
        c.alignment = Alignment(wrap_text=True, vertical="center")
        ws.column_dimensions[get_column_letter(i)].width = max(14, min(30, len(col) + 4))
    for i, col in enumerate(red.COLONNES_SYNTHESE, start=1):
        c = ws.cell(row=5, column=i, value=ligne[col])
        c.font = Font(size=9.5); c.border = _BORDER
        c.alignment = Alignment(wrap_text=True, vertical="top")
    ws.freeze_panes = "A5"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
